from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # --- Title Section ---
    title = doc.add_heading('Event Report: 3-Day AI Tools Workshop', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('Organized By: Computer Society Seniors')
    doc.add_paragraph('Dates: November 17 – November 19')
    
    doc.add_heading('1. Executive Summary', level=1)
    doc.add_paragraph(
        "The Computer Society Seniors successfully organized a three-day workshop aimed at "
        "demystifying Artificial Intelligence tools. The event bridged the gap between theoretical "
        "AI knowledge and practical application, focusing on productivity, creativity, and technical workflows."
    )

    # --- Day 1 ---
    doc.add_heading('2. Daily Proceedings', level=1)
    
    doc.add_heading('Day 1 (Nov 17): Foundations & Text Generation', level=2)
    doc.add_paragraph("Theme: Unlocking the Power of Large Language Models (LLMs)")
    
    p = doc.add_paragraph()
    p.add_run("Key Activities:").bold = True
    doc.add_paragraph('• Introduction to Neural Networks and Prompt Engineering.', style='List Bullet')
    doc.add_paragraph('• Hands-on "Prompt Battle" using ChatGPT and Gemini.', style='List Bullet')
    doc.add_paragraph('• Techniques for Zero-shot and Chain-of-thought prompting.', style='List Bullet')

    # --- Day 2 ---
    doc.add_heading('Day 2 (Nov 18): Visual Creativity & Multimedia', level=2)
    doc.add_paragraph("Theme: From Text to Pixel – AI in Design")
    
    p = doc.add_paragraph()
    p.add_run("Key Activities:").bold = True
    doc.add_paragraph('• Workshop on Midjourney and DALL-E 3 for brand assets.', style='List Bullet')
    doc.add_paragraph('• Ethical considerations in AI art and deepfakes.', style='List Bullet')
    doc.add_paragraph('• Challenge: Creating a full brand identity (Logo + Social Post).', style='List Bullet')

    # --- Day 3 (New Additions) ---
    doc.add_heading('Day 3 (Nov 19): Integration, Quiz & Presentations', level=2)
    doc.add_paragraph("Theme: Efficiency, Assessment, and Review")
    
    p = doc.add_paragraph()
    p.add_run("1. Workflow Automation & Coding:").bold = True
    doc.add_paragraph(
        "Sessions focused on using GitHub Copilot for coding and Zapier for workflow automation."
    )

    p = doc.add_paragraph()
    p.add_run("2. Student Presentations:").bold = True
    doc.add_paragraph(
        "Participants worked in groups to present their 'Brand Identity' projects created during Day 2. "
        "A panel of seniors reviewed the creative use of AI tools in their designs."
    )

    p = doc.add_paragraph()
    p.add_run("3. AI Tools Quiz:").bold = True
    doc.add_paragraph(
        "A live interactive quiz was conducted to test retention of concepts covered over the three days. "
        "Questions covered prompt engineering techniques, tool specific capabilities, and AI ethics."
    )

    p = doc.add_paragraph()
    p.add_run("4. Feedback Session:").bold = True
    doc.add_paragraph(
        "A structured feedback loop was established where attendees submitted anonymous responses "
        "regarding the workshop pace, content quality, and speaker effectiveness."
    )

    # --- Outcomes ---
    doc.add_heading('3. Key Outcomes', level=1)
    doc.add_paragraph('• 90% increased confidence in using AI tools.', style='List Bullet')
    doc.add_paragraph('• Creation of an "AI Toolkit PDF" shared with all members.', style='List Bullet')
    doc.add_paragraph('• Successful deployment of simple web apps by non-technical students.', style='List Bullet')

    # Save the file
    file_name = 'AI_Workshop_Report.docx'
    doc.save(file_name)
    print(f"File '{file_name}' has been created successfully!")

if __name__ == "__main__":
    create_report()